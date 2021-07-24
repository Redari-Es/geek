from docx import Document
from docx.shared import Inches

#创建对象
document = Document()
#
document.add_heading('Python爬虫', 0)
#
p = document.add_paragraph('Python爬虫开发-')
p.runs[0].bold = True
p.add_run('数据存储-').bold = True
#
p.add_run('Word-')
#添加内容并设置字体斜体
p.add_run('存储实例。').italic = True
#添加正文，设置“样式”-〉‘明显引用’
document.add_paragraph('样式' + '-' + '明显引用', style='IntenseQuote')
document.add_paragraph('项目符号1', style='ListBullet')
document.add_paragraph('项目符号2', style='ListNumber')
document.add_picture('pic.png', width=Inches(1.25))
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in range(2):
    row_cells = table.add_row().cells
    row_cells[0].text = 'a'
    row_cells[1].text = 'b'
    row_cells[2].text = 'c'
document.add_page_break()
document.save('test.docx')
